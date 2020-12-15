from tkinter import *
from tkinter import messagebox
from docx import Document
import openpyxl

from UI.TkinterGUI.page import Page
from Model.examiner import Examiner
from functools import partial

class ExaminerPage(Page):
    def __init__(self, *args, **kwargs):
        Page.__init__(self, *args, **kwargs)

        all = Frame(self)
        all.pack(side=LEFT, fill=Y)
        Label(all, text="All examiners").pack()
        self.all_listbox = Listbox(all, width=100)
        self.all_listbox.bind('<Double-1>', self.all_list_on_click)
        self.all_list = Examiner.all()
        for person in self.all_list:
            self.all_listbox.insert(END, person.to_string())
        self.all_listbox.pack(side="top", fill="both", expand=True)
        all_report = Button(all, text = "Export info about all examiners in docx files", command=self.all_docx_export)
        all_report.pack(side="top")
        xlsx_report = Button(all, text = "Export info about all examiners in Excel", command=self.xlsx_export)
        xlsx_report.pack(side="top")

        #################################################################################################

        find = Frame(self)
        find.pack(side=LEFT, fill=Y)
        Label(find, text="Find examiner").pack()
        find_entry = Entry(find)
        find_entry.pack()
        self.find_listbox = Listbox(find, width=100)
        self.find_listbox.bind('<Double-1>', self.find_list_on_click)
        find_examiner = partial(self.find_examiner, self.find_listbox, find_entry)
        find_btn = Button(find, text="Find", command = find_examiner).pack()
        self.find_listbox.pack(side="top", fill="both", expand=True)

        ########################################################################################################

        add = Frame(self)
        add.pack()

        Label(add, text="Add examiner", font="Helvetica -14 bold").grid(columnspan=2)

        Label(add, text="Surname: ").grid(row=1, column=0)
        surn = Entry(add, width=50)
        surn.grid(row=1, column=1)

        Label(add, text="Name: ").grid(row=2, column=0)
        name = Entry(add, width=50)
        name.grid(row=2, column=1)

        Label(add, text="Patronymic: ").grid(row=3, column=0)
        patr = Entry(add, width=50)
        patr.grid(row=3, column=1)

        Label(add, text="Salary: ").grid(row=4, column=0)
        salary = Entry(add, width=50)
        salary.grid(row=4, column=1)

        add_examiner = partial(self.add_examiner, surn, name, patr, salary, self.all_listbox)

        add_btn = Button(add, text="Add examiner", command =add_examiner)
        add_btn.grid(columnspan=2)

    def update_all(self, listbox):
        listbox.delete(0, END)
        self.all_list = Examiner.all()
        for elem in self.all_list:
            listbox.insert(END, elem.to_string())
        return

    def find_examiner(self, listbox, find_label):
        listbox.delete(0, END)
        self.find_list = Examiner.findBySurname(find_label.get())
        for elem in self.find_list:
            listbox.insert(END, elem.to_string())
        return

    def add_examiner(self, surn, name, patr, salary, listbox):
        Examiner.create(surn.get(), name.get(), patr.get(), int(salary.get()))
        surn.delete(0,END)
        name.delete(0, END)
        patr.delete(0, END)
        salary.delete(0, END)
        self.update_all(listbox)

    def all_list_on_click(self, event):
        cs = self.all_listbox.curselection()
        self.manipulateExaminer(self.all_list[cs[0]])

    def find_list_on_click(self, event):
        cs = self.find_listbox.curselection()
        self.manipulateExaminer(self.find_list[cs[0]])

    def manipulateExaminer(self, chosenExaminer):
        newWindow = Toplevel(self)
        newWindow.title("New Window")
        newWindow.geometry("1000x500")
        exams = Frame(newWindow)
        exams.pack(side=LEFT, fill=Y)
        Label(exams, text="Exams of examiner:").pack()
        exams_listbox = Listbox(exams, width=106)
        for person in chosenExaminer.get_exams():
            exams_listbox.insert(END, person.to_string())
        exams_listbox.pack(side=LEFT, fill="both", expand=True)
        change = Frame(newWindow)
        change.pack(side=RIGHT, fill="both")
        Label(change, text = "Change address").grid(row = 0, columnspan = 2)
        Label(change, text="Enter new address: ").grid(row = 1, column = 0)
        new_sal = Entry(change, width = "39")
        new_sal.delete(0, END)
        new_sal.insert(0, chosenExaminer.payment)
        new_sal.grid(row = 1, column = 1)
        change_sal = partial(self.change_salary, new_sal, chosenExaminer)
        change_sal_btn = Button(change, text = "Change salary", command=change_sal)
        change_sal_btn.grid(columnspan = 2)
        Label(change, text="------------------------------------------------------------").grid(columnspan = 2)
        delete = partial(self.delete_examiner, chosenExaminer)
        del_btn = Button(change, text = "Delete examiner", command=delete)
        del_btn.grid(columnspan = 2)
        Label(change, text="------------------------------------------------------------").grid(columnspan=2)
        examiner_export = partial(self.examiner_export, chosenExaminer)
        examiner_report = Button(change, text = "Export info about this examiner in Word", command=examiner_export)
        examiner_report.grid(columnspan = 2)

    def change_salary(self, sal_entry, chosenEnrollee):
        chosenEnrollee.change_payment(int(sal_entry.get()))
        self.update_all(self.all_listbox)

    def delete_examiner(self, chosenExaminer):
        result = messagebox.askquestion("Delete", "Are You Sure?", icon='warning')
        if result == 'yes':
            chosenExaminer.delete();
        else:
            return
        self.update_all(self.all_listbox)

    def xlsx_export(self):
        wb = openpyxl.load_workbook('..\\..\\Reports\\All.xlsx')
        if 'Examiners' not in wb.sheetnames:
            wb.create_sheet('Examiners')
        ws = wb.get_sheet_by_name('Examiners')
        ws.delete_cols(1, 6)
        ws.delete_rows(1, 100)
        for i in range (len(self.all_list)):
            ws.cell(row=i + 1, column=1).value = self.all_list[i].id
            ws.cell(row=i + 1, column=2).value = self.all_list[i].surname
            ws.cell(row=i + 1, column=3).value = self.all_list[i].name
            ws.cell(row=i + 1, column=4).value = self.all_list[i].patronymic
            ws.cell(row=i + 1, column=5).value = self.all_list[i].payment
        wb.save('..\\..\\Reports\\All.xlsx')

    def all_docx_export(self):
        for i in self.all_list:
            self.examiner_export(i)

    def examiner_export(self, examiner):
        document = Document()
        document.add_heading((examiner.surname + ' ' + examiner.name + ' ' + examiner.patronymic + "(Examiner)"), 0)
        document.add_heading('Overall information', level=1)
        document.add_paragraph('ID: ' + str(examiner.id))
        document.add_paragraph('Name: ' + examiner.surname + ' ' + examiner.name + ' ' + examiner.patronymic)
        document.add_paragraph('Salary: ' + str(examiner.payment))
        document.add_heading('Exams', level=1)
        table = document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Time'
        hdr_cells[3].text = 'Status'
        hdr_cells[4].text = 'Score'
        hdr_cells[5].text = 'Enrollee'
        for exam in examiner.get_exams():
            row_cells = table.add_row().cells
            row_cells[0].text = str(exam.id)
            row_cells[1].text = exam.exam_name
            row_cells[2].text = exam.pass_time
            row_cells[3].text = exam.status
            row_cells[4].text = str(exam.score)
            row_cells[5].text = exam.get_enrollee()

        document.save('D:\\Alex\\Work\\Study\\Programming\\Python\\Python_labs_3.1\\Reports\\Examiners\\'
                      + examiner.surname + '.' + examiner.name[0] + '.' + examiner.patronymic[0] + '.(' + str(examiner.id) + ').docx')