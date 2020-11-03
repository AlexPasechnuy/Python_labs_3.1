from tkinter import *
from tkinter import messagebox
from tkcalendar import DateEntry
from docx import Document
import openpyxl

from UI.GUI.page import Page
from Model.enrollee import Enrollee
from functools import partial


class EnrolleePage(Page):
    def __init__(self, *args, **kwargs):
        Page.__init__(self, *args, **kwargs)

        all = Frame(self)
        all.pack(side=LEFT, fill=Y)
        Label(all, text="All enrollees").pack()
        self.all_listbox = Listbox(all, width=100)
        self.all_listbox.bind('<Double-1>', self.all_list_on_click)
        self.all_list = Enrollee.all()
        for person in self.all_list:
            self.all_listbox.insert(END, person.to_string())
        self.all_listbox.pack(side="top", fill="both", expand=True)
        all_report = Button(all, text = "Export info about all enrollees in Excel", command=self.all_export)
        all_report.pack(side="top")

        #################################################################################################

        find = Frame(self)
        find.pack(side=LEFT, fill=Y)
        Label(find, text="Find enrollee").pack()
        find_entry = Entry(find)
        find_entry.pack()
        self.find_listbox = Listbox(find, width=100)
        self.find_listbox.bind('<Double-1>', self.find_list_on_click)
        find_enr = partial(self.find_enr, self.find_listbox, find_entry)
        find_btn = Button(find, text="Find", command = find_enr).pack()
        self.find_listbox.pack(side="top", fill="both", expand=True)

        ########################################################################################################

        add = Frame(self)
        add.pack()

        Label(add, text="Add enrollee", font="Helvetica -14 bold").grid(columnspan=2)

        Label(add, text="Surname: ").grid(row=1, column=0)
        surn = Entry(add, width=50)
        surn.grid(row=1, column=1)

        Label(add, text="Name: ").grid(row=2, column=0)
        name = Entry(add, width=50)
        name.grid(row=2, column=1)

        Label(add, text="Patronymic: ").grid(row=3, column=0)
        patr = Entry(add, width=50)
        patr.grid(row=3, column=1)

        Label(add, text="Address: ").grid(row=4, column=0)
        addr = Entry(add, width=50)
        addr.grid(row=4, column=1)

        Label(add, text="Date of birthday: ").grid(row=5, column=0)
        birth = DateEntry(add, width=12, background='darkblue', foreground='white', borderwidth=2, date_pattern='dd.mm.y')
        birth.grid(row=5, column=1)

        Label(add, text="Number of passport: ").grid(row=6, column=0)
        passp = Entry(add, width=50)
        passp.grid(row=6, column=1)

        add_enr = partial(self.add_enr, surn, name, patr, addr, birth, passp, self.all_listbox)

        add_btn = Button(add, text="Add enrollee", command =add_enr)
        add_btn.grid(columnspan=2)

    def update_all(self, listbox):
        listbox.delete(0, END)
        self.all_list = Enrollee.all()
        for elem in self.all_list:
            listbox.insert(END, elem.to_string())
        return

    def find_enr(self, listbox, find_label):
        listbox.delete(0, END)
        self.find_list = Enrollee.findBySurname(find_label.get())
        for elem in self.find_list:
            listbox.insert(END, elem.to_string())
        return

    def add_enr(self,surn, name, patr, addr, birth, passp, listbox):
        Enrollee.create(surn.get(), name.get(), patr.get(), addr.get(), birth.get(), passp.get())
        surn.delete(0,END)
        name.delete(0, END)
        patr.delete(0, END)
        addr.delete(0, END)
        birth.delete(0, END)
        passp.delete(0, END)
        self.update_all(listbox)

    def all_list_on_click(self, event):
        cs = self.all_listbox.curselection()
        self.manipulateEnrollee(self.all_list[cs[0]])

    def find_list_on_click(self, event):
        cs = self.find_listbox.curselection()
        self.manipulateEnrollee(self.find_list[cs[0]])

    def manipulateEnrollee(self, chosenEnrollee):
        newWindow = Toplevel(self)
        newWindow.title("New Window")
        newWindow.geometry("1000x500")
        exams = Frame(newWindow)
        exams.pack(side=LEFT, fill=Y)
        Label(exams, text="Exams of enrollee:").pack()
        exams_listbox = Listbox(exams, width=106)
        for person in chosenEnrollee.get_exams():
            exams_listbox.insert(END, person.to_string())
        exams_listbox.pack(side=LEFT, fill="both", expand=True)
        change = Frame(newWindow)
        change.pack(side=RIGHT, fill="both")
        Label(change, text = "Change address").grid(row = 0, columnspan = 2)
        Label(change, text="Enter new address: ").grid(row = 1, column = 0)
        new_addr = Entry(change, width = "39")
        new_addr.delete(0,END)
        new_addr.insert(0,chosenEnrollee.address)
        new_addr.grid(row = 1, column = 1)
        change_addr = partial(self.change_addr, new_addr, chosenEnrollee)
        change_addr_btn = Button(change, text = "Change address", command=change_addr)
        change_addr_btn.grid(columnspan = 2)
        Label(change, text="------------------------------------------------------------").grid(columnspan = 2)
        delete = partial(self.delete_enr, chosenEnrollee)
        del_btn = Button(change, text = "Delete enrollee", command=delete)
        del_btn.grid(columnspan = 2)
        Label(change, text="------------------------------------------------------------").grid(columnspan=2)
        enr_export = partial(self.enr_export, chosenEnrollee)
        enr_report = Button(change, text = "Export info about this enrollee in Word", command=enr_export)
        enr_report.grid(columnspan = 2)

    def change_addr(self, addr_entry, chosenEnrollee):
        chosenEnrollee.change_address(addr_entry.get())
        self.update_all(self.all_listbox)

    def delete_enr(self, chosenEnrollee):
        result = messagebox.askquestion("Delete", "Are You Sure?", icon='warning')
        if result == 'yes':
            chosenEnrollee.delete();
        else:
            return
        self.update_all(self.all_listbox)

    def all_export(self):
        wb = openpyxl.load_workbook('D:\\Alex\\Work\\Study\\Programming\\Python\\Python_labs_3.1\\Reports\\All.xlsx')
        if 'Enrollees' not in wb.sheetnames:
            wb.create_sheet('Enrollees')
        ws = wb.get_sheet_by_name('Enrollees')
        ws.delete_cols(1, 6)
        ws.delete_rows(1, 100)
        for i in range (len(self.all_list)):
            ws.cell(row=i + 1, column=1).value = self.all_list[i].id
            ws.cell(row=i + 1, column=2).value = self.all_list[i].surname
            ws.cell(row=i + 1, column=3).value = self.all_list[i].name
            ws.cell(row=i + 1, column=4).value = self.all_list[i].patronymic
            ws.cell(row=i + 1, column=5).value = self.all_list[i].address
            ws.cell(row=i + 1, column=6).value = self.all_list[i].birthday
            ws.cell(row=i + 1, column=7).value = self.all_list[i].passport

        wb.save('..\\..\\Reports\\All.xlsx')

    def enr_export(self, enr):
        document = Document()

        document.add_heading((enr.surname + ' ' + enr.name + ' ' + enr.patronymic + "(Enrollee)"), 0)
        document.add_heading('Overall information', level=1)
        document.add_paragraph('ID: ' + str(enr.id))
        document.add_paragraph('Name: ' + enr.surname + ' ' + enr.name + ' ' + enr.patronymic)
        document.add_paragraph('Address: ' + enr.address)
        document.add_paragraph('Birthday: ' + enr.birthday)
        document.add_paragraph('Number of passport: ' + enr.passport)
        document.add_heading('Exams', level=1)
        for i in enr.get_exams():
            document.add_paragraph(
                i.to_string(), style='List Bullet'
            )

        document.save('..\\..\\Reports\\Enrollees\\'
                      + enr.surname + '.' + enr.name[0] + '.' + enr.patronymic[0] + '.(' + str(enr.id) + ').docx')
