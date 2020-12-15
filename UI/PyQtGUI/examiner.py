from PyQt5 import QtWidgets, uic
import sys
import openpyxl
from docx import Document

from PyQt5.QtGui import QStandardItem, QStandardItemModel
from PyQt5.QtWidgets import QAbstractItemView, QMainWindow, QApplication, QPushButton, QVBoxLayout, QLabel, QWidget

from Model.examiner import Examiner

def examiner_export(examiner):
    document = Document()
    document.add_heading((examiner.surname + ' ' + examiner.name + ' ' + examiner.patronymic + "(Examiner)"), 0)
    document.add_heading('Overall information', level=1)
    document.add_paragraph('ID: ' + str(examiner.id))
    document.add_paragraph('Name: ' + examiner.surname + ' ' + examiner.name + ' ' + examiner.patronymic)
    document.add_paragraph('Salary: ' + str(examiner.payment))
    document.add_heading('Exams', level=1)
    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Time'
    hdr_cells[3].text = 'Status'
    hdr_cells[4].text = 'Score'
    hdr_cells[5].text = 'Enrollee'
    for exam in examiner.get_exams():
        row_cells = table.add_row().cells
        row_cells[0].text = str(exam.id)
        row_cells[1].text = exam.exam_name
        row_cells[2].text = exam.pass_time
        row_cells[3].text = exam.status
        row_cells[4].text = str(exam.score)
        row_cells[5].text = exam.get_enrollee()

    document.save('D:\\Alex\\Work\\Study\\Programming\\Python\\Python_labs_3.1\\Reports\\Examiners\\'
                  + examiner.surname + '.' + examiner.name[0] + '.' + examiner.patronymic[0] + '.(' + str(
        examiner.id) + ').docx')

class ManipulateExaminerWindow(QMainWindow):
    def __init__(self, id):
        super().__init__()
        uic.loadUi("layouts\\manipulateExaminer.ui", self)
        self.examiner = Examiner.get_by_id(id)
        self.newSal.setValue(self.examiner.payment)
        self.changeSal.clicked.connect(self.changeSalaryClick)
        self.deleteExaminer.clicked.connect(self.deleteExaminerClick)
        self.exportBtn.clicked.connect(self.exportExaminerClick)
        self.updateAllList()

    def updateAllList(self):
        self.allList.setEditTriggers(QAbstractItemView.NoEditTriggers)
        model = QStandardItemModel(self.allList)
        for i in self.examiner.get_exams():
            item = QStandardItem(i.to_string())
            model.appendRow(item)
        self.allList.setModel(model)

    def changeSalaryClick(self):
        self.examiner.change_payment(self.newSal.value())

    def deleteExaminerClick(self):
        self.examiner.delete()
        self.close()

    def exportExaminerClick(self):
        examiner_export(self.examiner)



class ExaminerWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi("layouts\\examiner.ui", self)
        self.findExaminerButton.clicked.connect(self.findExaminerOnClick)
        self.addExaminerButton.clicked.connect(self.addExaminerOnClick)
        self.manipulateExaminerButton.clicked.connect(self.editExaminerClick)
        self.updateAllListBtn.clicked.connect(self.updateAllList)
        self.exportDocx.clicked.connect(self.exportDocxClick)
        self.exportXlsx.clicked.connect(self.exportXlsxClick)
        self.updateAllList()

    def updateAllList(self):
        self.allExaminersList.setEditTriggers(QAbstractItemView.NoEditTriggers)
        model = QStandardItemModel(self.allExaminersList)
        for i in Examiner.all():
            item = QStandardItem(i.to_string())
            model.appendRow(item)
        self.allExaminersList.setModel(model)

    def findExaminerOnClick(self):
        model = QStandardItemModel(self.findExaminerList)
        for i in Examiner.findBySurname(self.findSurnameInput.text()):
            item = QStandardItem(i.to_string())
            model.appendRow(item)
        self.findExaminerList.setModel(model)

    def addExaminerOnClick(self):
        Examiner.create(self.surnameEdit.text(), self.nameEdit.text(),
                        self.patronymicEdit.text(), self.salaryEdit.value())
        self.updateAllList()

    def editExaminerClick(self):
        id = int(self.examinerIDInput.value())
        self.w = ManipulateExaminerWindow(id)
        self.w.show()

    def exportDocxClick(self):
        for i in Examiner.all():
            examiner_export(i)

    def exportXlsxClick(self):
        wb = openpyxl.load_workbook('..\\..\\Reports\\All.xlsx')
        if 'Examiners' not in wb.sheetnames:
            wb.create_sheet('Examiners')
        ws = wb['Examiners']
        ws.delete_cols(1, 6)
        ws.delete_rows(1, 100)
        all_list = Examiner.all()
        for i in range (len(all_list)):
            ws.cell(row=i + 1, column=1).value = all_list[i].id
            ws.cell(row=i + 1, column=2).value = all_list[i].surname
            ws.cell(row=i + 1, column=3).value = all_list[i].name
            ws.cell(row=i + 1, column=4).value = all_list[i].patronymic
            ws.cell(row=i + 1, column=5).value = all_list[i].payment
        wb.save('..\\..\\Reports\\All.xlsx')

# app = QApplication(sys.argv)
# w = ExaminerWindow()
# w.show()
# app.exec_()