from PyQt5 import QtWidgets, uic, QtCore
import sys
import openpyxl
from docx import Document

from PyQt5.QtGui import QStandardItem, QStandardItemModel
from PyQt5.QtWidgets import QAbstractItemView, QMainWindow, QApplication, QPushButton, QVBoxLayout, QLabel, QWidget

from Model.enrollee import Enrollee
from Model.examiner import Examiner
from Model.exam import Exam


def exam_export(exam):
    document = Document()

    document.add_heading((exam.exam_name + "(Exam)"), 0)
    document.add_heading('Overall information', level=1)
    document.add_paragraph('ID: ' + str(exam.id))
    document.add_paragraph('Name: ' + exam.exam_name)
    document.add_paragraph('Pass time: ' + exam.pass_time)
    document.add_paragraph('Status: ' + exam.status)
    document.add_paragraph('Score: ' + str(exam.score))
    document.add_paragraph('Enrollee: ' + str(exam.get_enrollee()))
    document.add_paragraph('Examiner: ' + str(exam.get_examiner()))

    document.save('..\\..\\Reports\\Exams\\'
                  + exam.exam_name + '.(' + str(exam.id) + ').docx')


class ManipulateExamWindow(QMainWindow):
    def __init__(self, id):
        super().__init__()
        uic.loadUi("layouts\\manipulateExam.ui", self)
        self.exam = Exam.get_by_id(id)
        self.newTime.setDateTime(QtCore.QDateTime.currentDateTime())
        self.changeTimeButton.clicked.connect(self.changeTimeClick)
        self.deleteExam.clicked.connect(self.deleteExamClick)
        self.exportBtn.clicked.connect(self.exportExamClick)

    def changeTimeClick(self):
        self.exam.change_time(self.newTime.dateTime().toString("dd.MM.yyyy hh:mm"))

    def deleteExamClick(self):
        self.exam.delete()
        self.close()

    def exportExamClick(self):
        exam_export(self.exam)



class ExamWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi("layouts\\exam.ui", self)
        self.findExamButton.clicked.connect(self.findExamOnClick)
        self.addExamBtn.clicked.connect(self.addExamClick)
        self.manipulateExamButton.clicked.connect(self.editExamClick)
        self.updateAllListBtn.clicked.connect(self.updateAllList)
        self.exportDocx.clicked.connect(self.exportDocxClick)
        self.exportXlsx.clicked.connect(self.exportXlsxClick)
        self.findEnrolleeButton.clicked.connect(self.findEnrolleeClick)
        self.findExaminerButton.clicked.connect(self.findExaminerClick)
        self.updateAllList()

    def updateAllList(self):
        self.allExamsList.setEditTriggers(QAbstractItemView.NoEditTriggers)
        model = QStandardItemModel(self.allExamsList)
        for i in Exam.all():
            item = QStandardItem(i.to_string())
            model.appendRow(item)
        self.allExamsList.setModel(model)

    def findExamOnClick(self):
        model = QStandardItemModel(self.findExamList)
        for i in Exam.find_by_date(self.findExamInput.date().toString("dd.MM.yyyy")):
            item = QStandardItem(i.to_string())
            model.appendRow(item)
        self.findExamList.setModel(model)

    def findEnrolleeClick(self):
        model = QStandardItemModel(self.findEnrolleeList)
        for i in Enrollee.findBySurname(self.findEnrolleeInput.text()):
            item = QStandardItem(i.to_string())
            model.appendRow(item)
        self.findEnrolleeList.setModel(model)

    def findExaminerClick(self):
        model = QStandardItemModel(self.findExaminerList)
        for i in Examiner.findBySurname(self.findExaminerInput.text()):
            item = QStandardItem(i.to_string())
            model.appendRow(item)
        self.findExaminerList.setModel(model)

    def addExamClick(self):
        Exam.create(self.passTimeInput.dateTime().toString("dd.MM.yyyy hh:mm"),
                        self.nameInput.text(),
                        self.scoreInput.value(), self.enrolleeID.value(),
                        self.examinerID.value())
        self.updateAllList()

    def editExamClick(self):
        id = int(self.examID.value())
        self.w = ManipulateExamWindow(id)
        self.w.show()

    def exportDocxClick(self):
        for i in Exam.all():
            exam_export(i)

    def exportXlsxClick(self):
        wb = openpyxl.load_workbook('..\\..\\Reports\\All.xlsx')
        if 'Exams' not in wb.sheetnames:
            wb.create_sheet('Exams')
        ws = wb['Exams']
        ws.delete_cols(1, 6)
        ws.delete_rows(1, 100)
        all_list = Exam.all()
        for i in range (len(all_list)):
            ws.cell(row=i + 1, column=1).value = all_list[i].id
            ws.cell(row=i + 1, column=2).value = all_list[i].exam_name
            ws.cell(row=i + 1, column=3).value = all_list[i].pass_time
            ws.cell(row=i + 1, column=4).value = all_list[i].status
            ws.cell(row=i + 1, column=5).value = all_list[i].score
            ws.cell(row=i + 1, column=6).value = all_list[i].get_enrollee()
            ws.cell(row=i + 1, column=7).value = all_list[i].get_examiner()
        wb.save('..\\..\\Reports\\All.xlsx')

# app = QApplication(sys.argv)
# w = ExamWindow()
# w.show()
# app.exec_()