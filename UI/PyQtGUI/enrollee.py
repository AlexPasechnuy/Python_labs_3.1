from PyQt5 import QtWidgets, uic
import sys
import openpyxl
from docx import Document
from PyQt5.QtGui import QStandardItem, QStandardItemModel
from PyQt5.QtWidgets import QAbstractItemView, QMainWindow, QApplication, QPushButton, QVBoxLayout, QLabel, QWidget

from Model.enrollee import Enrollee

from UI.PyQtGUI.examiner import ExaminerWindow
from UI.PyQtGUI.exam import ExamWindow

def enr_export(enr):
    document = Document()

    document.add_heading((enr.surname + ' ' + enr.name + ' ' + enr.patronymic + "(Enrollee)"), 0)
    document.add_heading('Overall information', level=1)
    document.add_paragraph('ID: ' + str(enr.id))
    document.add_paragraph('Name: ' + enr.surname + ' ' + enr.name + ' ' + enr.patronymic)
    document.add_paragraph('Address: ' + enr.address)
    document.add_paragraph('Birthday: ' + enr.birthday)
    document.add_paragraph('Number of passport: ' + enr.passport)
    document.add_heading('Exams', level=1)
    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Time'
    hdr_cells[3].text = 'Status'
    hdr_cells[4].text = 'Score'
    hdr_cells[5].text = 'Examiner'
    for exam in enr.get_exams():
        row_cells = table.add_row().cells
        row_cells[0].text = str(exam.id)
        row_cells[1].text = exam.exam_name
        row_cells[2].text = exam.pass_time
        row_cells[3].text = exam.status
        row_cells[4].text = str(exam.score)
        row_cells[5].text = exam.get_examiner()

    document.save('..\\..\\Reports\\Enrollees\\'
                  + enr.surname + '.' + enr.name[0] + '.' + enr.patronymic[0] + '.(' + str(enr.id) + ').docx')


class ManipulateEnrWindow(QMainWindow):
    def __init__(self, id):
        super().__init__()
        uic.loadUi("layouts\\manipulateEnrollee.ui", self)
        self.enrollee = Enrollee.get_by_id(id)
        self.newAddr.setText(self.enrollee.address)
        self.changeAddr.clicked.connect(self.changeAddressClick)
        self.deleteEnr.clicked.connect(self.deleteEnrolleeClick)
        self.exportBtn.clicked.connect(self.exportEnrolleeClick)
        self.updateAllList()

    def updateAllList(self):
        self.allList.setEditTriggers(QAbstractItemView.NoEditTriggers)
        model = QStandardItemModel(self.allList)
        for i in self.enrollee.get_exams():
            item = QStandardItem(i.to_string())
            model.appendRow(item)
        self.allList.setModel(model)

    def changeAddressClick(self):
        self.enrollee.change_address(self.newAddr.text())

    def deleteEnrolleeClick(self):
        self.enrollee.delete()
        self.close()

    def exportEnrolleeClick(self):
        enr_export(self.enrollee)

class EnrolleeWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi("layouts\\enrollee.ui", self)
        self.findEnrolleeButton.clicked.connect(self.findEnrolleeOnClick)
        self.addEnrBtn.clicked.connect(self.addEnrolleeOnClick)
        self.manipulateEnrButton.clicked.connect(self.editEnrClick)
        self.updateAllListBtn.clicked.connect(self.updateAllList)
        self.exportDocx.clicked.connect(self.exportDocxClick)
        self.exportXlsx.clicked.connect(self.exportXlsxClick)
        self.goToExaminerBtn.clicked.connect(self.goToExaminerClick)
        self.goToExamBtn.clicked.connect(self.goToExamClick)
        self.updateAllList()

    def updateAllList(self):
        self.allEnrList.setEditTriggers(QAbstractItemView.NoEditTriggers)
        model = QStandardItemModel(self.allEnrList)
        for i in Enrollee.all():
            item = QStandardItem(i.to_string())
            model.appendRow(item)
        self.allEnrList.setModel(model)

    def findEnrolleeOnClick(self):
        model = QStandardItemModel(self.findEnrList)
        for i in Enrollee.findBySurname(self.findSurnameInput.text()):
            item = QStandardItem(i.to_string())
            model.appendRow(item)
        self.findEnrList.setModel(model)

    def addEnrolleeOnClick(self):
        Enrollee.create(self.surnameEdit.text(), self.nameEdit.text(),
                        self.patronymicEdit.text(), self.addressEdit.text(),
                        self.birthdayEdit.date().toString("dd.MM.yyyy"), self.passportEdit.text())
        self.updateAllList()

    def editEnrClick(self):
        id = int(self.enrIDInput.value())
        self.w = ManipulateEnrWindow(id)
        self.w.show()

    def exportDocxClick(self):
        for i in Enrollee.all():
            enr_export(i)

    def exportXlsxClick(self):
        wb = openpyxl.load_workbook('..\\..\\Reports\\All.xlsx')
        if 'Enrollees' not in wb.sheetnames:
            wb.create_sheet('Enrollees')
        ws = wb['Enrollees']
        ws.delete_cols(1, 6)
        ws.delete_rows(1, 100)
        all_list = Enrollee.all()
        for i in range (len(all_list)):
            ws.cell(row=i + 1, column=1).value = all_list[i].id
            ws.cell(row=i + 1, column=2).value = all_list[i].surname
            ws.cell(row=i + 1, column=3).value = all_list[i].name
            ws.cell(row=i + 1, column=4).value = all_list[i].patronymic
            ws.cell(row=i + 1, column=5).value = all_list[i].address
            ws.cell(row=i + 1, column=6).value = all_list[i].birthday
            ws.cell(row=i + 1, column=7).value = all_list[i].passport
        wb.save('..\\..\\Reports\\All.xlsx')

    def goToExaminerClick(self):
        self.w1 = ExaminerWindow()
        self.w1.show()

    def goToExamClick(self):
        self.w2 = ExamWindow()
        self.w2.show()

app = QApplication(sys.argv)
w = EnrolleeWindow()
w.show()
app.exec_()